# -*- coding: utf-8 -*-
import zipfile, shutil, os
from fpdf import FPDF
from docx import Document

DOCX = "/Users/gintaresarmaviciute/Downloads/Builder APPlication/Oliver_Varcoe_CV.docx"
OUT  = "/Users/gintaresarmaviciute/Downloads/Builder APPlication/Oliver_Varcoe_CV.pdf"
PHOTO_TMP = "/tmp/cv_photo_tmp.png"

REPLACEMENTS = [
    ("–", "-"), ("—", "-"), ("•", "-"),
    ("‘", "'"), ("’", "'"), ("“", '"'), ("”", '"'),
    ("·", "*"), (" ", " "),
    ("é", "e"), ("è", "e"), ("ê", "e"),
    ("à", "a"), ("â", "a"), ("ä", "a"),
    ("î", "i"), ("ï", "i"),
    ("ô", "o"), ("ö", "o"),
    ("ù", "u"), ("û", "u"), ("ü", "u"),
    ("ç", "c"), ("æ", "ae"), ("ø", "o"),
    ("É", "E"), ("À", "A"), ("Ç", "C"),
    ("Ö", "O"), ("Ü", "U"),
]

def clean(t):
    for k, v in REPLACEMENTS:
        t = t.replace(k, v)
    return t.encode("latin-1", errors="replace").decode("latin-1")

def extract_photo():
    with zipfile.ZipFile(DOCX) as z:
        imgs = [n for n in z.namelist() if n.startswith("word/media/") and n != "word/media/"]
        if not imgs:
            return None
        with z.open(imgs[0]) as src, open(PHOTO_TMP, "wb") as dst:
            shutil.copyfileobj(src, dst)
    return PHOTO_TMP

def first_run_bold(p):
    for r in p.runs:
        if r.text.strip():
            return bool(r.bold)
    return False


class CV(FPDF):
    def header(self): pass

    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "", 8)
        self.set_text_color(160, 155, 150)
        self.cell(0, 5, str(self.page_no()), align="C")

    def section_heading(self, title):
        self.ln(4)
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(30, 25, 23)
        self.set_x(self.l_margin)
        self.cell(0, 6, clean(title).upper())
        self.ln(6)
        self.set_draw_color(180, 180, 180)
        self.set_line_width(0.3)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def role_line(self, bold_part, normal_part, indent=0):
        usable = self.w - self.l_margin - self.r_margin - indent
        self.set_font("Helvetica", "B", 9.5)
        name_w = self.get_string_width(clean(bold_part)) + 1
        self.set_font("Helvetica", "", 9)
        date_w = self.get_string_width(clean(normal_part)) + 1

        self.set_x(self.l_margin + indent)
        self.set_font("Helvetica", "B", 9.5)
        self.set_text_color(30, 25, 23)
        self.cell(name_w, 5.5, clean(bold_part))
        filler = usable - name_w - date_w
        if filler > 0:
            self.cell(filler, 5.5, "")
        self.set_font("Helvetica", "", 9)
        self.set_text_color(100, 95, 90)
        self.cell(date_w, 5.5, clean(normal_part), align="R")
        self.ln(5.5)
        self.set_text_color(30, 25, 23)

    def company_line(self, runs):
        bold_text = ""
        normal_text = ""
        hit_normal = False
        for r in runs:
            if not r.text:
                continue
            if r.bold and not hit_normal:
                bold_text += r.text
            else:
                hit_normal = True
                normal_text += r.text
        normal_text = normal_text.strip().lstrip("\t").strip()
        self.role_line(bold_text.strip(), normal_text)

    def sub_role_line(self, runs):
        bold_text = ""
        normal_text = ""
        hit_normal = False
        for r in runs:
            if not r.text:
                continue
            if r.bold and not hit_normal:
                bold_text += r.text
            else:
                hit_normal = True
                normal_text += r.text
        normal_text = normal_text.strip().lstrip("\t").strip()
        self.role_line(bold_text.strip(), normal_text, indent=4)

    def body_text(self, text, indent=0):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(80, 75, 70)
        if indent:
            self.set_x(self.l_margin + indent)
        self.multi_cell(0, 4.5, clean(text))
        self.set_text_color(30, 25, 23)
        self.ln(1)

    def bullet(self, text, indent=4):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(60, 55, 50)
        self.set_x(self.l_margin + indent)
        self.cell(4, 4.5, "-")
        self.multi_cell(0, 4.5, clean(text))
        self.set_text_color(30, 25, 23)

    def lang_line(self, runs):
        self.set_font("Helvetica", "", 9)
        self.set_text_color(60, 55, 50)
        x = self.l_margin
        for r in runs:
            if not r.text:
                continue
            if r.bold:
                self.set_font("Helvetica", "B", 9)
                self.set_text_color(30, 25, 23)
            else:
                self.set_font("Helvetica", "", 9)
                self.set_text_color(80, 75, 70)
            self.set_x(x)
            w = self.get_string_width(clean(r.text))
            self.cell(w, 5, clean(r.text))
            x += w
        self.ln(5)
        self.set_text_color(30, 25, 23)


def build():
    doc = Document(DOCX)
    pdf = CV("P", "mm", "A4")
    pdf.set_margins(18, 18, 18)
    pdf.set_auto_page_break(True, 16)
    pdf.add_page()

    # Header
    photo_path = extract_photo()
    photo_size = 28

    hdr_cell = doc.tables[0].rows[0].cells[0].text.strip()
    lines = [l for l in hdr_cell.split("\n") if l.strip()]
    name_line = lines[0] if lines else ""
    subtitle_line = lines[1] if len(lines) > 1 else ""
    contact_lines = lines[2:] if len(lines) > 2 else []

    content_w = pdf.w - pdf.l_margin - pdf.r_margin - photo_size - 6

    if photo_path and os.path.exists(photo_path):
        pdf.image(photo_path, x=pdf.w - pdf.r_margin - photo_size,
                  y=pdf.t_margin, w=photo_size, h=photo_size)

    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(30, 25, 23)
    pdf.set_x(pdf.l_margin)
    pdf.cell(content_w, 8, clean(name_line))
    pdf.ln(8)

    if subtitle_line:
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(100, 95, 90)
        pdf.set_x(pdf.l_margin)
        pdf.cell(content_w, 5, clean(subtitle_line))
        pdf.ln(5)
        pdf.ln(2)

    pdf.set_font("Helvetica", "", 8.5)
    pdf.set_text_color(80, 75, 70)
    for cl in contact_lines:
        pdf.set_x(pdf.l_margin)
        pdf.cell(content_w, 4.5, clean(cl))
        pdf.ln(4.5)

    sep_y = max(pdf.get_y() + 4, pdf.t_margin + photo_size + 4)
    pdf.set_y(sep_y)
    pdf.set_draw_color(180, 180, 180)
    pdf.set_line_width(0.3)
    pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
    pdf.ln(5)

    # Body
    paras = doc.paragraphs
    n = len(paras)
    i = 1  # skip blank first para

    while i < n:
        p = paras[i]
        text = p.text.strip()
        sname = p.style.name if p.style else ""

        if not text:
            i += 1
            continue

        if sname == "Heading 2":
            pdf.section_heading(text)
            i += 1
            continue

        if sname == "List Paragraph":
            pdf.bullet(text)
            i += 1
            continue

        if first_run_bold(p):
            # Look ahead for context
            next_i = i + 1
            while next_i < n and not paras[next_i].text.strip():
                next_i += 1
            next_is_list = (next_i < n and paras[next_i].style and
                            paras[next_i].style.name == "List Paragraph")

            is_sub_role = ("\t" in p.text and next_is_list and
                           "·" not in p.text and "·" not in p.text)

            if is_sub_role:
                pdf.sub_role_line(p.runs)
            elif "\t" not in p.text and not next_is_list:
                # Languages or education with mixed bold/normal runs
                pdf.lang_line(p.runs)
            else:
                pdf.company_line(p.runs)
                pdf.ln(0.5)
        else:
            pdf.body_text(text)

        i += 1

    pdf.output(OUT)
    print("Saved: {}  ({:,} bytes)".format(OUT, os.path.getsize(OUT)))

build()
